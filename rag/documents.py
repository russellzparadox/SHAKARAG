"""Document ingestion for SHAKARAG.

Extracts text from PDF, Word (.docx), and Excel (.xlsx/.xls) files, chunks it,
and stores chunks in the SAME vector collection as the schema index — tagged
with kind="document" so retrieval can mix schema knowledge and document knowledge,
and the router can tell them apart by metadata.
"""
from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from .chunking import Chunk

DOC_CHUNK_CHARS = 1400
DOC_CHUNK_OVERLAP = 180
MAX_FILE_BYTES = 40 * 1024 * 1024  # 40 MB guard


class DocIngestError(Exception):
    pass


def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def _clean(text: str) -> str:
    text = text.replace("\u00ad", "")           # soft hyphens
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(data: bytes) -> list[str]:
    """Return list of page texts."""
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise DocIngestError("pymupdf is not installed") from exc
    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pages.append(_clean(page.get_text("text")))
    return pages


def extract_docx(data: bytes) -> list[str]:
    """Return paragraph + table-cell text as one block (split later)."""
    try:
        import io

        from docx import Document
    except ImportError as exc:
        raise DocIngestError("python-docx is not installed") from exc
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            style = getattr(para.style, "name", "") or ""
            if style.lower().startswith("heading"):
                level = "".join(ch for ch in style if ch.isdigit()) or "1"
                parts.append("#" * int(level) + " " + t)
            else:
                parts.append(t)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))
    return [_clean("\n".join(parts))]


def extract_xlsx(data: bytes) -> list[str]:
    """One text block per sheet; rows rendered as pipe-delimited lines with a header row."""
    try:
        import io

        import openpyxl
    except ImportError as exc:
        raise DocIngestError("openpyxl is not installed") from exc
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets: list[str] = []
    for ws in wb.worksheets:
        lines: list[str] = [f"## Sheet: {ws.title}"]
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v).strip() for v in row]
            if any(cells):
                lines.append(" | ".join(cells))
            if len(lines) > 4000:  # per-sheet guard
                lines.append("(truncated)")
                break
        if len(lines) > 1:
            sheets.append(_clean("\n".join(lines)))
    wb.close()
    return sheets


EXTRACTORS = {
    ".pdf": ("pdf", extract_pdf),
    ".docx": ("word", extract_docx),
    ".xlsx": ("excel", extract_xlsx),
    ".xlsm": ("excel", extract_xlsx),
}


def extract_any(filename: str, data: bytes) -> tuple[str, list[str]]:
    """Dispatch on extension. Returns (kind, sections[])."""
    suffix = Path(filename).suffix.lower()
    if suffix not in EXTRACTORS:
        raise DocIngestError(
            f"Unsupported file type '{suffix}'. Supported: pdf, docx, xlsx, xlsm."
        )
    if len(data) > MAX_FILE_BYTES:
        raise DocIngestError("File too large (max 40 MB).")
    if not data:
        raise DocIngestError("Empty file.")
    kind, extractor = EXTRACTORS[suffix]
    sections = extractor(data)
    sections = [s for s in sections if s]
    if not sections:
        raise DocIngestError(
            "No extractable text found — the file may be scanned images or empty."
        )
    return kind, sections


def chunk_sections(
    sections: list[str],
    source_name: str,
    kind: str,
    doc_id: str | None = None,
) -> list[Chunk]:
    """Split sections into overlapping chunks with document metadata."""
    doc_id = doc_id or _sha((source_name + sections[0][:200]).encode())
    chunks: list[Chunk] = []

    # Excel sheets become pseudo-table cards so the router can treat them
    # like queryable tables.
    if kind == "excel":
        for section in sections:
            title_line = next(
                (l for l in section.splitlines() if l.startswith("## Sheet:")), ""
            )
            sheet_name = title_line.replace("## Sheet:", "").strip() or source_name
            text = f"SPREADSHEET: {source_name} — {title_line.lstrip('# ')}\n\n{section}"
            chunks.extend(_split_text(text, source_name, kind, doc_id, sheet_name))
        return chunks

    for i, section in enumerate(sections, start=1):
        label = f"page {i}" if kind == "pdf" and len(sections) > 1 else ""
        chunks.extend(_split_text(section, source_name, kind, doc_id, label))
    return chunks


def _split_text(
    text: str, source_name: str, kind: str, doc_id: str, section_label: str
) -> list[Chunk]:
    if len(text) <= DOC_CHUNK_CHARS:
        return [_mk_chunk(text, source_name, kind, doc_id, section_label, 0)]
    step = DOC_CHUNK_CHARS - DOC_CHUNK_OVERLAP
    out: list[Chunk] = []
    for i, start in enumerate(range(0, len(text), step)):
        piece = text[start : start + DOC_CHUNK_CHARS]
        out.append(_mk_chunk(piece, source_name, kind, doc_id, section_label, i))
        if start + DOC_CHUNK_CHARS >= len(text):
            break
    return out


def _mk_chunk(
    text: str, source_name: str, kind: str, doc_id: str, section_label: str, part: int
) -> Chunk:
    header = f"[DOCUMENT {kind.upper()}] Source: {source_name}"
    if section_label:
        header += f" — {section_label}"
    cid = f"doc:{doc_id}:{part}:{abs(hash(section_label)) % 99999}"
    return Chunk(
        id=cid,
        text=f"{header}\n{text}",
        metadata={
            "kind": "document",
            "doc_kind": kind,
            "source": source_name,
            "doc_id": doc_id,
            "section": section_label,
            "part": part,
        },
    )


def ingest_document_bytes(
    store: Any, filename: str, data: bytes
) -> dict[str, Any]:
    """Extract → chunk → upsert into the vector store. Returns stats."""
    kind, sections = extract_any(filename, data)
    chunks = chunk_sections(sections, filename, kind)
    n = store.upsert(chunks)
    return {
        "filename": filename,
        "kind": kind,
        "sections": len(sections),
        "chunks": len(chunks),
        "upserted": n,
        "store_count": store.count,
    }
